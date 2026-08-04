"""
写作相关 API 路由 — 需登录。
"""
import json
import base64
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse

from ..schemas import (
    GenerateFrameworkRequest, GenerateContentRequest, RefineRequest, FrameworkResponse,
)
from ..services.writer_service import (
    generate_framework, generate_content_stream, refine_document,
    save_document, detect_doc_type,
)
from ..utils.auth import get_current_user_id

router = APIRouter(prefix="/api/writing", tags=["写作"])


@router.post("/generate-framework")
async def api_generate_framework(
    req: GenerateFrameworkRequest,
    user_id: str = Depends(get_current_user_id),
):
    """生成公文框架/提纲（流式，避免超时）"""
    try:
        result = await generate_framework(
            topic=req.topic,
            doc_type=req.doc_type,
            keywords=req.keywords,
        )
        return {
            "title_suggestion": result["title_suggestion"],
            "framework": result["framework"],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"框架生成失败：{str(e)}")


@router.post("/detect-doc-type")
async def api_detect_doc_type(
    topic: str = Form(...),
    user_id: str = Depends(get_current_user_id),
):
    """根据主题自动识别文种"""
    try:
        result = await detect_doc_type(topic)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文种识别失败：{str(e)}")


@router.post("/generate-content")
async def api_generate_content(
    req: GenerateContentRequest,
    user_id: str = Depends(get_current_user_id),
):
    """流式生成公文内容（SSE）"""
    async def event_stream():
        async for event in generate_content_stream(
            topic=req.topic,
            doc_type=req.doc_type,
            keywords=req.keywords,
            framework=req.framework,
            custom_instructions=req.custom_instructions,
            reference_material=req.custom_instructions or "",
        ):
            event_type = event["type"]
            data = event["data"]

            if event_type == "error":
                yield f"event: error\ndata: {json.dumps({'error': str(data)}, ensure_ascii=False)}\n\n"
                return

            if event_type == "complete":
                try:
                    result = data
                    doc_id = save_document(
                        user_id=user_id,
                        title=result.get("title", req.topic),
                        doc_type=req.doc_type,
                        keywords=req.keywords,
                        framework=result.get("framework", req.framework or []),
                        content=result.get("content", ""),
                    )
                    result["document_id"] = doc_id
                except Exception:
                    result = data
                    result["document_id"] = ""
                yield f"event: complete\ndata: {json.dumps(result, ensure_ascii=False)}\n\n"
                return

            yield f"event: {event_type}\ndata: {json.dumps({'text': str(data)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/refine")
async def api_refine(
    req: RefineRequest,
    user_id: str = Depends(get_current_user_id),
):
    """润色已有文稿"""
    try:
        result = await refine_document(
            content=req.content,
            doc_type=req.doc_type,
            instructions=req.instructions or "",
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"润色失败：{str(e)}")


@router.post("/upload-reference")
async def api_upload_reference(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user_id),
):
    """上传参考材料，返回提取的文本内容"""
    try:
        content = await file.read()
        # 尝试提取文本
        text = ""
        filename = file.filename or ""

        if filename.lower().endswith(('.txt', '.md')):
            text = content.decode('utf-8', errors='ignore')
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')):
            # 图片：base64 编码传给 Claude 视觉能力
            b64 = base64.b64encode(content).decode()
            ext = filename.split('.')[-1].lower() or 'png'
            text = f"[图片文件: {filename}]\n用户上传了一张参考图片，请根据图片内容提取关键信息。\n"
            text += f"图片数据(base64): data:image/{ext};base64,{b64[:500]}...（截断）"
        elif filename.lower().endswith('.pdf'):
            text = f"[PDF文件: {filename}]\nPDF内容无法直接提取，请手动将关键内容粘贴到写作主题中。"
        elif filename.lower().endswith(('.doc', '.docx')):
            text = f"[Word文档: {filename}]\n文档内容无法直接提取，请手动将关键内容粘贴到写作主题中。"
        else:
            text = content.decode('utf-8', errors='ignore')

        return {
            "filename": filename,
            "text": text[:5000],  # 限制长度
            "length": len(text),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件处理失败：{str(e)}")


# ========== 审校与导出 ==========

from ..services.audit_service import audit_document
from fastapi.responses import Response


@router.post("/audit")
async def api_audit(content: str = Form(...), user_id: str = Depends(get_current_user_id)):
    """审校文稿——错别字/格式/敏感词/逻辑 四项检查"""
    try:
        result = await audit_document(content)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"审校失败：{str(e)}")


@router.post("/export-word")
async def api_export_word(content: str = Form(...), title: str = Form("公文"), user_id: str = Depends(get_current_user_id)):
    """导出为Word文档（GB/T 9704-2012标准格式）"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        import io

        doc = DocxDocument()

        # A4页面设置
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)

        # 标题 —— 二号方正小标宋（回退用黑体）
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 正文 —— 三号仿宋
        paragraphs = content.replace('\r\n', '\n').replace('\r', '\n').split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()

            # 判断是否为标题行
            is_heading = False
            for prefix in ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、',
                          '（一）', '（二）', '（三）', '（四）', '（五）',
                          '1.', '2.', '3.', '4.', '5.']:
                if para_text.startswith(prefix):
                    is_heading = True
                    break

            if para_text.startswith('#'):
                is_heading = True
                para_text = para_text.lstrip('#').strip()

            run = p.add_run(para_text)
            run.font.size = Pt(16)
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

            if is_heading:
                run.font.bold = True
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            # 首行缩进2字符
            p.paragraph_format.first_line_indent = Pt(32)
            p.paragraph_format.line_spacing = Pt(28)

        # 保存到内存
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"{title}.docx"
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败：{str(e)}")
